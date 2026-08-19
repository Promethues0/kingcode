#!/usr/bin/env python3
"""用 openpyxl / python-docx 独立回读校验 KingCode 生成的 xlsx/docx。

自己生成、自己解析等于自证。这个脚本用完全独立的第三方实现回读，
才算证明「Excel / Word 真能打开」。

依赖（仅校验用，运行时不需要）：pip install openpyxl python-docx
跑法：python3 test/verify-ooxml.py
"""
import subprocess
import sys
import tempfile
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
failed = 0


def check(ok, label, extra=""):
    global failed
    if not ok:
        failed += 1
    print(f"{'OK  ' if ok else 'FAIL'} {label}{('  ' + extra) if extra else ''}")


def main():
    try:
        import openpyxl
        import docx
    except ImportError as e:
        print(f"跳过：缺少校验依赖（{e}）。pip install openpyxl python-docx")
        return 0

    with tempfile.TemporaryDirectory() as tmp:
        xlsx_path = Path(tmp) / "t.xlsx"
        docx_path = Path(tmp) / "t.docx"
        script = f"""
import('{REPO}/presets/mpe-assess/lib/ooxml.js').then(async m => {{
  const fs = await import('node:fs');
  fs.writeFileSync({str(xlsx_path)!r}, m.buildXlsx([
    {{ name: '差距矩阵', columns: ['指标','判定','得分'], rows: [
      ['通信机密性','不符合', 0.25],
      ['存储机密性 <特殊&字符>','未测评', null],
    ]}},
    {{ name: '非法/名:称*', columns: ['a'], rows: [['b']] }},
  ]));
  fs.writeFileSync({str(docx_path)!r}, m.buildDocx([
    {{ type:'heading', level:1, text:'密评自评估报告' }},
    {{ type:'paragraph', text:'第一行\\n第二行' }},
    {{ type:'table', columns:['项目','数量'], rows:[['可判定','3']] }},
  ]));
}})
"""
        r = subprocess.run(["node", "-e", script], capture_output=True, text=True)
        if r.returncode != 0:
            print("生成失败：", r.stderr[:400])
            return 1

        wb = openpyxl.load_workbook(xlsx_path)
        check(wb.sheetnames[0] == "差距矩阵", "xlsx：工作表名", str(wb.sheetnames))
        check("/" not in wb.sheetnames[1] and ":" not in wb.sheetnames[1],
              "xlsx：非法字符已净化", wb.sheetnames[1])
        ws = wb["差距矩阵"]
        check(ws["A1"].font.bold is True, "xlsx：表头加粗")
        check(ws.freeze_panes == "A2", "xlsx：首行冻结", str(ws.freeze_panes))
        check(ws.auto_filter.ref is not None, "xlsx：自动筛选")
        check(isinstance(ws["C2"].value, float) and ws["C2"].value == 0.25,
              "xlsx：数字保持数值类型", repr(ws["C2"].value))
        check(ws["A3"].value == "存储机密性 <特殊&字符>", "xlsx：XML 特殊字符往返", repr(ws["A3"].value))
        check(ws["C3"].value is None, "xlsx：null 单元格为空")

        d = docx.Document(docx_path)
        styles = [p.style.name for p in d.paragraphs if p.text.strip()]
        check(any(s.startswith("Heading 1") for s in styles), "docx：一级标题样式", str(styles))
        body = "\n".join(p.text for p in d.paragraphs)
        check("第一行" in body and "第二行" in body, "docx：段落换行保留")
        check(len(d.tables) == 1 and d.tables[0].rows[0].cells[0].text == "项目",
              "docx：表格与表头")

    print(f"\n{'全部通过' if failed == 0 else f'失败 {failed} 项'}")
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
